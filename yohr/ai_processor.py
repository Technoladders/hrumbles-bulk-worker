"""
yohr/ai_processor.py
Stage 3 — extract full text from PDF/DOCX, send to GPT for structured extraction.

Key changes vs previous version:
  - Full SYSTEM_PROMPT: extracts work_experience, education, projects, certifications
  - Text COMPRESSION before AI call — removes redundant whitespace to save input tokens
  - FULL resume text stored (no 3000-char truncation)
  - DOCX support via python-docx (fallback to raw byte scan for .doc)
  - MAX_AI_TOKENS = 4096 (was 700 — too small for full structured output)
  - Null-byte sanitization retained
"""
import io
import json
import logging
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

from .constants import (
    supabase, openai_client, YOHR_ORG_ID,
    STORAGE_BUCKET,
    OPENAI_MODEL, MAX_AI_TOKENS, MAX_AI_INPUT_CHARS,
    MAX_AI_WORKERS, MAX_AI_RETRIES,
)

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Prompt — parse exactly what is in the resume, no generation
# ---------------------------------------------------------------------------
SYSTEM_PROMPT = """
Based on the provided resume text, perform a detailed extraction to create a professional profile.
Parse exactly what is written — do NOT generate, infer, or summarize beyond what is in the text.
Return ONLY a single, valid JSON object. No markdown. No explanations.

Fields (use null or [] if not found):
"suggested_title": string — most recent job title exactly as written in the resume
"candidate_name": string — full name
"email": string — email address
"phone": string — phone exactly as written
"linkedin_url": string or null — must be a valid URL
"github_url": string or null — must be a valid URL
"current_location": string or null — city/address
"professional_summary": array of strings — exact text from Summary/Profile/Overview section split by sentences or bullets
"top_skills": array of strings — all skills listed exactly as written
"work_experience": array of {"company": str, "designation": str, "duration": str, "responsibilities": [str]}
"education": array of {"institution": str, "degree": str, "year": str}
"projects": array of strings — copy project blocks fully, do not summarize
"certifications": array of strings
"other_details": object or null — any other structured sections (languages, awards, publications, etc.)
"total_experience": string or null — total experience exactly as stated in resume (e.g. "5 years", "12+ years")
"current_company": string or null — most recent employer name
"current_designation": string or null — most recent title exactly as written
"notice_period": string or null — if mentioned
"highest_education": string or null — highest degree and institution on one line
""".strip()

EMPTY_AI_RESULT: dict = {
    "suggested_title": None,
    "candidate_name": None,
    "email": None,
    "phone": None,
    "linkedin_url": None,
    "github_url": None,
    "current_location": None,
    "professional_summary": [],
    "top_skills": [],
    "work_experience": [],
    "education": [],
    "projects": [],
    "certifications": [],
    "other_details": None,
    "total_experience": None,
    "current_company": None,
    "current_designation": None,
    "notice_period": None,
    "highest_education": None,
}


# ---------------------------------------------------------------------------
# Text utilities
# ---------------------------------------------------------------------------

def _sanitize(text: str) -> str:
    """Strip null bytes PostgreSQL rejects (error 22P05)."""
    return text.replace('\x00', '').replace('\u0000', '')


def _compress(text: str) -> str:
    """
    Compress whitespace to reduce AI input token count.
    Preserves ALL content — only removes redundant formatting characters.
    Typical saving: 15-30% fewer tokens.
    """
    # Null bytes
    text = _sanitize(text)
    # Common unicode noise
    text = text.replace('\xa0', ' ').replace('\u2003', ' ').replace('\ufeff', '')
    # Tabs → single space
    text = re.sub(r'\t+', ' ', text)
    # Multiple spaces → one
    text = re.sub(r'  +', ' ', text)
    # 3+ newlines → 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Strip each line
    lines = [line.strip() for line in text.split('\n')]
    # Drop lines that are purely decorative (dashes, dots, underscores, stars)
    lines = [l for l in lines if not re.fullmatch(r'[-_=.*•·▪▸►●○|+]{2,}', l)]
    return '\n'.join(l for l in lines if l).strip()


# ---------------------------------------------------------------------------
# Text extraction — PDF and DOCX
# ---------------------------------------------------------------------------

def _extract_pdf(pdf_bytes: bytes) -> str:
    """Extract text from PDF bytes using pypdf with pdfminer fallback."""
    # Attempt 1: pypdf (fast)
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
        pages = [p.extract_text() or "" for p in reader.pages]
        text = "\n".join(pages).strip()
        if text:
            return text
    except Exception as exc:
        logger.debug("pypdf failed: %s", exc)

    # Attempt 2: pdfminer.six (better layout handling)
    try:
        from pdfminer.high_level import extract_text as pm_extract
        return pm_extract(io.BytesIO(pdf_bytes)) or ""
    except Exception as exc:
        logger.debug("pdfminer failed: %s", exc)

    return ""


def _extract_docx(docx_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        import docx as python_docx
        doc = python_docx.Document(io.BytesIO(docx_bytes))
        parts: list[str] = []

        # Paragraphs (main body text)
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)

        # Tables (skills tables, education tables, etc.)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)
    except Exception as exc:
        logger.warning("python-docx extraction failed: %s", exc)
        return ""


def _extract_doc_raw(doc_bytes: bytes) -> str:
    """
    Best-effort text extraction from legacy .doc files.
    Scans for readable ASCII/UTF-8 runs — imperfect but better than nothing.
    """
    try:
        # Try interpreting as UTF-16 (some .doc files)
        text = doc_bytes.decode("utf-16-le", errors="ignore")
        # Keep only printable chars + newlines
        text = re.sub(r'[^\x20-\x7E\n\t]', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()
        if len(text) > 100:
            return text
    except Exception:
        pass

    # Fallback: raw ASCII extraction
    text = doc_bytes.decode("latin-1", errors="replace")
    text = re.sub(r'[^\x20-\x7E\n]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text if len(text) > 100 else ""


def _extract_text(storage_path: str) -> str:
    """
    Download file from Supabase storage and extract text.
    Supports: .pdf, .docx, .doc (best-effort)
    """
    file_bytes: bytes = supabase.storage.from_(STORAGE_BUCKET).download(storage_path)
    ext = Path(storage_path).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(file_bytes)
    elif ext == ".docx":
        return _extract_docx(file_bytes)
    elif ext == ".doc":
        return _extract_doc_raw(file_bytes)
    else:
        # Try PDF first (some files have wrong extension), then docx
        text = _extract_pdf(file_bytes)
        if not text:
            text = _extract_docx(file_bytes)
        return text


# ---------------------------------------------------------------------------
# AI call
# ---------------------------------------------------------------------------

def _call_ai(full_text: str) -> dict:
    """Send compressed resume text to GPT and parse structured JSON response."""
    compressed = _compress(full_text)
    ai_input = compressed[:MAX_AI_INPUT_CHARS]

    response = openai_client.chat.completions.create(
        model=OPENAI_MODEL,
        max_tokens=MAX_AI_TOKENS,
        temperature=0,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user",   "content": ai_input},
        ],
    )
    raw = (response.choices[0].message.content or "").strip()

    # Strip markdown fences if present
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    raw = raw.strip()

    try:
        parsed = json.loads(raw)
    except json.JSONDecodeError:
        # Try to detect truncation (incomplete JSON) and surface a clear error
        preview = raw[:200].replace('\n', ' ')
        raise ValueError(f"AI returned non-JSON or truncated response: {preview}")

    # Merge with defaults so callers can safely .get() any key
    result = {**EMPTY_AI_RESULT, **parsed}

    # Normalise list fields
    for list_key in ("professional_summary", "top_skills", "work_experience",
                     "education", "projects", "certifications"):
        if not isinstance(result.get(list_key), list):
            result[list_key] = []

    return result


# ---------------------------------------------------------------------------
# Row processor
# ---------------------------------------------------------------------------

def _process_row(row: dict) -> None:
    row_id   = row["id"]
    attempts = (row.get("s3_attempts") or 0) + 1

    supabase.table("org_csv_import_rows").update(
        {"s3_status": "processing", "s3_attempts": attempts}
    ).eq("id", row_id).execute()

    try:
        # ── Extract full resume text ─────────────────────────────────────
        full_text = row.get("resume_text_excerpt") or ""
        if not full_text and row.get("stored_resume_path"):
            full_text = _extract_text(row["stored_resume_path"])

        if not full_text:
            # Build minimal context from CSV fields so AI still returns structure
            parts = []
            for k, label in [("raw_name", "Name"), ("raw_designation", "Title"),
                              ("raw_company", "Company"), ("raw_location", "Location")]:
                if row.get(k):
                    parts.append(f"{label}: {row[k]}")
            full_text = "\n".join(parts) or f"Candidate: {row.get('raw_name', 'Unknown')}"

        # Sanitize before storing
        full_text = _sanitize(full_text)

        # ── Call AI ───────────────────────────────────────────────────────
        ai_result = _call_ai(full_text)

        # ── Persist ───────────────────────────────────────────────────────
        supabase.table("org_csv_import_rows").update({
            "s3_status":           "done",
            "s3_error":            None,
            "resume_text_excerpt": full_text,   # store FULL text
            "ai_result":           ai_result,
        }).eq("id", row_id).execute()

        top_skills = ai_result.get("top_skills") or []
        logger.debug("ai_processor: row %s done — %d skills, %d jobs",
                     row_id, len(top_skills), len(ai_result.get("work_experience") or []))

    except Exception as exc:
        error_msg  = _sanitize(str(exc))
        new_status = "failed" if attempts >= MAX_AI_RETRIES else "pending"
        logger.warning("ai_processor: row %s attempt %d failed: %s", row_id, attempts, error_msg)
        supabase.table("org_csv_import_rows").update({
            "s3_status":   new_status,
            "s3_attempts": attempts,
            "s3_error":    error_msg,
        }).eq("id", row_id).execute()


# ---------------------------------------------------------------------------
# Scheduler entry point
# ---------------------------------------------------------------------------

def run_ai_processor() -> None:
    try:
        rows = (
            supabase.table("org_csv_import_rows")
            .select(
                "id, session_id, stored_resume_path, raw_name, raw_designation, "
                "raw_company, raw_location, s3_attempts, resume_text_excerpt"
            )
            .eq("org_id", YOHR_ORG_ID)
            .in_("s2_status", ["done", "skipped"])
            .eq("s3_status", "pending")
            .limit(50)
            .execute()
            .data
        )
    except Exception as exc:
        logger.error("ai_processor: fetch failed: %s", exc)
        return

    if not rows:
        return

    logger.info("ai_processor: processing %d rows", len(rows))

    with ThreadPoolExecutor(max_workers=MAX_AI_WORKERS) as pool:
        futures = {pool.submit(_process_row, row): row for row in rows}
        for future in as_completed(futures):
            row = futures[future]
            try:
                future.result()
            except Exception as exc:
                logger.error("ai_processor: unhandled error for row %s: %s", row["id"], exc)

    session_ids = {r["session_id"] for r in rows}
    for sid in session_ids:
        try:
            supabase.rpc("refresh_csv_session_counts", {"p_session_id": sid}).execute()
        except Exception as exc:
            logger.warning("ai_processor: refresh_counts failed for %s: %s", sid, exc)