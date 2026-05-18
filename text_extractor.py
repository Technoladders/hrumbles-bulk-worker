"""
text_extractor.py — PDF / DOCX / DOC text extraction
"""

import io
import os
import logging
import requests
from typing import Tuple

logger = logging.getLogger(__name__)
MIN_TEXT_LENGTH = 100


def _extract_pdfminer(file_bytes: bytes) -> str:
    from pdfminer.high_level import extract_text_to_fp
    from pdfminer.layout import LAParams
    output = io.StringIO()
    extract_text_to_fp(io.BytesIO(file_bytes), output,
                       laparams=LAParams(detect_vertical=False, all_texts=True))
    return output.getvalue()


def _extract_pypdf(file_bytes: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    parts = [page.extract_text() for page in reader.pages if page.extract_text()]
    return '\n'.join(parts)


def _extract_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return '\n'.join(parts)


def _extract_doc_antiword(file_bytes: bytes) -> str:
    import subprocess, tempfile
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as f:
            f.write(file_bytes)
            tmp_path = f.name
        result = subprocess.run(
            ['antiword', '-w', '0', tmp_path],
            capture_output=True, text=True, timeout=30,
            encoding='utf-8', errors='replace',
        )
        return result.stdout or ''
    except FileNotFoundError:
        logger.error('antiword not found — add to Dockerfile')
        return ''
    except subprocess.TimeoutExpired:
        logger.error('antiword timed out')
        return ''
    except Exception as e:
        logger.error(f'antiword error: {e}')
        return ''
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except Exception:
                pass


def _extract_doc_via_parser(storage_path: str) -> str:
    parser_url = os.getenv('RESUME_PARSER_URL', 'http://resume-parser-container:5005')
    try:
        resp = requests.post(f'{parser_url}/api/extract-doc-text',
                             json={'storage_path': storage_path}, timeout=120)
        if resp.status_code == 200:
            return resp.json().get('text', '')
    except Exception as e:
        logger.warning(f'Parser container fallback failed: {e}')
    return ''


def extract_text(file_bytes: bytes, mime_type: str, storage_path: str = None) -> Tuple[str, str]:
    mime = (mime_type or '').lower()
    name = (storage_path or '').lower()

    # ── PDF ───────────────────────────────────────────────────────────────────
    if 'pdf' in mime or name.endswith('.pdf'):
        try:
            text = _extract_pdfminer(file_bytes)
            if len(text.strip()) >= MIN_TEXT_LENGTH:
                return text.strip(), 'pdfminer'
        except Exception as e:
            logger.warning(f'pdfminer failed: {e}')
        try:
            text = _extract_pypdf(file_bytes)
            if len(text.strip()) >= MIN_TEXT_LENGTH:
                return text.strip(), 'pypdf'
        except Exception as e:
            logger.warning(f'pypdf failed: {e}')
        return '', 'image_only'

    # ── DOCX ──────────────────────────────────────────────────────────────────
    elif (
        'wordprocessingml' in mime
        or 'docx' in mime
        or name.endswith('.docx')
        or mime == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    ):
        try:
            text = _extract_docx(file_bytes)
            if len(text.strip()) >= 50:
                return text.strip(), 'docx'
            return '', 'failed'
        except Exception as e:
            logger.error(f'DOCX failed: {e}')
            return '', 'failed'

    # ── DOC ───────────────────────────────────────────────────────────────────
    elif mime == 'application/msword' or 'msword' in mime or name.endswith('.doc'):
        # Strategy 1: antiword
        text = _extract_doc_antiword(file_bytes)
        if len(text.strip()) >= 50:
            return text.strip(), 'antiword'
        # Strategy 2: existing parser container
        if storage_path:
            text = _extract_doc_via_parser(storage_path)
            if len(text.strip()) >= 50:
                return text.strip(), 'doc_api'
        logger.warning(f'DOC failed completely: {storage_path}')
        return '', 'failed'

    # ── Unsupported ───────────────────────────────────────────────────────────
    else:
        logger.warning(f'Unsupported: {mime_type}')
        return '', 'unsupported'